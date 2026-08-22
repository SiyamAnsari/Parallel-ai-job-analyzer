"""
Resume parser supporting PDF, DOCX, TXT, and raw markdown files.
Extracts sections, skills, years of experience, and structured candidate profile.
"""
import io
import re
from typing import Optional, List, Dict, Union, BinaryIO
from pathlib import Path
from src.schemas.models import CandidateProfile

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None


class ResumeParser:
    """Production-grade resume parser extracting text and structure from multiple formats."""

    @staticmethod
    def extract_text_from_file(file_source: Union[str, Path, BinaryIO, bytes], filename: Optional[str] = None) -> str:
        """Extract text from a file path, file-like object, or raw bytes."""
        if isinstance(file_source, (str, Path)):
            path = Path(file_source)
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                return ResumeParser._extract_from_pdf(path.read_bytes())
            elif suffix in [".docx", ".doc"]:
                return ResumeParser._extract_from_docx(path.read_bytes())
            else:
                return path.read_text(encoding="utf-8", errors="ignore")
        elif isinstance(file_source, bytes):
            ext = (filename or "").lower()
            if ext.endswith(".pdf") or (len(file_source) > 4 and file_source.startswith(b"%PDF")):
                return ResumeParser._extract_from_pdf(file_source)
            elif ext.endswith(".docx"):
                return ResumeParser._extract_from_docx(file_source)
            else:
                return file_source.decode("utf-8", errors="ignore")
        elif hasattr(file_source, "read"):
            content = file_source.read()
            return ResumeParser.extract_text_from_file(content, filename=filename)
        else:
            raise ValueError(f"Unsupported file source type: {type(file_source)}")

    @staticmethod
    def _extract_from_pdf(pdf_bytes: bytes) -> str:
        """Extract plain text from PDF bytes using pypdf."""
        if PdfReader is None:
            raise ImportError("pypdf is required to parse PDF resumes. Please install it.")
        
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()

    @staticmethod
    def _extract_from_docx(docx_bytes: bytes) -> str:
        """Extract plain text from DOCX bytes using python-docx."""
        if docx is None:
            raise ImportError("python-docx is required to parse DOCX resumes. Please install it.")
        
        doc = docx.Document(io.BytesIO(docx_bytes))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()

    @classmethod
    def parse_to_profile(cls, raw_text_or_file: Union[str, Path, bytes, BinaryIO], filename: Optional[str] = None, candidate_name: Optional[str] = None) -> CandidateProfile:
        """Parse raw text or resume file into a CandidateProfile model."""
        if isinstance(raw_text_or_file, str) and not raw_text_or_file.endswith((".pdf", ".docx", ".doc", ".txt")):
            raw_text = raw_text_or_file.strip()
        else:
            raw_text = cls.extract_text_from_file(raw_text_or_file, filename=filename).strip()

        # Fast deterministic heuristics
        extracted_name = candidate_name or cls._extract_candidate_name(raw_text)
        years_exp = cls._estimate_years_of_experience(raw_text)
        skills = cls._extract_skills(raw_text)
        highlights = cls._extract_experience_highlights(raw_text)
        education = cls._extract_education(raw_text)

        return CandidateProfile(
            name=extracted_name,
            target_title=cls._extract_target_title(raw_text),
            years_of_experience=years_exp,
            skills=skills,
            experience_highlights=highlights,
            education=education,
            raw_text=raw_text
        )

    @staticmethod
    def _extract_candidate_name(text: str) -> str:
        """Extract first line or name pattern as candidate name."""
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if not lines:
            return "Candidate"
        
        # Check first 3 lines
        for line in lines[:3]:
            # If line is short and looks like a name (not an email, URL, or section title)
            if len(line) < 40 and not any(char in line for char in ["@", "http", "www", ":", "RESUME", "CURRICULUM"]):
                # Clean special characters
                clean_name = re.sub(r"[^a-zA-Z\s\.\-]", "", line).strip()
                if 2 <= len(clean_name.split()) <= 4:
                    return clean_name
        return lines[0][:40]

    @staticmethod
    def _estimate_years_of_experience(text: str) -> float:
        """Find mentions like '3+ years', '5 years of experience', or calculate date ranges."""
        pattern = r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs|year)\s*(?:of\s*)?(?:experience|exp)?"
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            try:
                numbers = [float(m) for m in matches if float(m) < 40]
                if numbers:
                    return max(numbers)
            except ValueError:
                pass
        
        # Check year spans (e.g. 2019 - 2024)
        year_ranges = re.findall(r"\b(20\d{2}|19\d{2})\s*(?:-|–|to)\s*(20\d{2}|present|current)\b", text, re.IGNORECASE)
        if year_ranges:
            total_years = 0
            for start_yr, end_yr in year_ranges:
                try:
                    s = int(start_yr)
                    e = 2026 if end_yr.lower() in ["present", "current"] else int(end_yr)
                    if 0 <= (e - s) <= 40:
                        total_years = max(total_years, e - s)
                except ValueError:
                    continue
            if total_years > 0:
                return float(total_years)

        return 0.0

    @staticmethod
    def _extract_skills(text: str) -> List[str]:
        """Extract prominent tech and analytical skills using keyword catalog and patterns."""
        known_skill_catalog = [
            "python", "sql", "excel", "power bi", "tableau", "r", "pandas", "numpy", "scikit-learn",
            "tensorflow", "pytorch", "aws", "azure", "gcp", "docker", "kubernetes", "git", "ci/cd",
            "fastapi", "django", "flask", "react", "node.js", "javascript", "typescript", "html", "css",
            "java", "c++", "c#", "scala", "spark", "hadoop", "kafka", "snowflake", "bigquery",
            "dbt", "airflow", "data analysis", "data engineering", "machine learning", "deep learning",
            "nlp", "computer vision", "llm", "langchain", "prompt engineering", "statistics",
            "business intelligence", "etl", "postgresql", "mysql", "mongodb", "redis", "linux",
            "agile", "scrum", "jira", "communication", "leadership", "problem solving"
        ]
        
        found_skills = set()
        lower_text = text.lower()
        
        for skill in known_skill_catalog:
            # Word boundary matching
            pattern = rf"\b{re.escape(skill)}\b"
            if re.search(pattern, lower_text):
                found_skills.add(skill.title() if len(skill) > 3 else skill.upper())
        
        # Also look under 'Skills:' or 'Technical Skills:' sections
        skills_section = re.search(r"(?:skills|technical skills|technologies)[\s:]*\n([\s\S]*?)(?:\n\n|\n[A-Z][A-Za-z\s]+:|$)", text, re.IGNORECASE)
        if skills_section:
            items = re.split(r"[,•|\n;]", skills_section.group(1))
            for item in items:
                cleaned = item.strip().strip("•-* ")
                if 2 <= len(cleaned) <= 30 and not any(k in cleaned.lower() for k in ["experience", "education", "project"]):
                    found_skills.add(cleaned)
                    
        return sorted(list(found_skills))[:25]

    @staticmethod
    def _extract_experience_highlights(text: str) -> List[str]:
        """Extract bullet points or major achievement sentences."""
        bullets = re.findall(r"(?:^|\n)\s*[•\-\*]\s*(.+)", text)
        if bullets:
            return [b.strip() for b in bullets if len(b.strip()) > 15][:10]
        
        # Fallback to sentences with action verbs
        sentences = [s.strip() for s in re.split(r"[\n\.]+", text) if len(s.strip()) > 25]
        return sentences[:8]

    @staticmethod
    def _extract_education(text: str) -> List[str]:
        """Extract degrees, diplomas, and certifications."""
        edu_keywords = ["bachelor", "master", "phd", "b.s.", "m.s.", "b.sc", "m.sc", "b.tech", "degree", "certified", "certification", "university", "institute", "college"]
        results = []
        for line in text.split("\n"):
            line_str = line.strip()
            if any(k in line_str.lower() for k in edu_keywords) and len(line_str) < 120:
                results.append(line_str.strip("•-* "))
        return results[:5]

    @staticmethod
    def _extract_target_title(text: str) -> Optional[str]:
        """Extract likely job title target."""
        titles = [
            "Data Analyst", "Senior Data Analyst", "Data Scientist", "Data Engineer",
            "Machine Learning Engineer", "AI Engineer", "Software Engineer",
            "Full Stack Developer", "Backend Developer", "Frontend Developer",
            "Product Manager", "DevOps Engineer", "Cloud Architect", "Business Analyst"
        ]
        for t in titles:
            if re.search(rf"\b{re.escape(t)}\b", text, re.IGNORECASE):
                return t
        return None

